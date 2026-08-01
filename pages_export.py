"""
pages_export.py — Export Center: TXT, Markdown, JSON, PDF and DOCX downloads,
built on the original reportlab / python-docx logic, with premium formatting.
"""

import json
from io import BytesIO
from datetime import datetime

import streamlit as st
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor
from docx import Document
from docx.shared import Pt, RGBColor

import db
from styles import glass_open, glass_close
from components import page_header, empty_state


def _get_active_project_and_script():
    pid = st.session_state.get("active_project_id")
    project = db.get_project(pid) if pid else None
    script_text = st.session_state.get("current_script_text") or (project["script"] if project else None)
    return project, script_text


def _build_pdf(title, script_text):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, title=title or "Scriptoria Export")
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ScriptoriaTitle", parent=styles["Title"], textColor=HexColor("#8a6d1a"), spaceAfter=18,
    )
    body_style = ParagraphStyle("ScriptoriaBody", parent=styles["Normal"], leading=15)

    elements = [Paragraph(title or "Untitled Project", title_style), Spacer(1, 12)]
    for line in script_text.split("\n"):
        safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        elements.append(Paragraph(safe_line if safe_line.strip() else "&nbsp;", body_style))
        elements.append(Spacer(1, 4))
    doc.build(elements)
    buffer.seek(0)
    return buffer


def _build_docx(title, script_text):
    buffer = BytesIO()
    document = Document()
    heading = document.add_heading(title or "Untitled Project", level=0)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x8A, 0x6D, 0x1A)
    for line in script_text.split("\n"):
        p = document.add_paragraph(line)
        p.style.font.size = Pt(11)
    document.save(buffer)
    buffer.seek(0)
    return buffer


def render():
    page_header("📤 Export Center", "Download your project in any format")

    project, script_text = _get_active_project_and_script()

    if not script_text:
        empty_state("Generate a script first, then come back here to export it.")
        return

    title = project["title"] if project else st.session_state.get("current_project_title", "Untitled Project")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")

    glass_open()
    st.write(f"Exporting **{title}**")

    col1, col2, col3, col4, col5 = st.columns(5)

    project_id = project["id"] if project else None

    with col1:
        if st.download_button("📄 TXT", script_text, file_name=f"{title}_{timestamp}.txt", use_container_width=True):
            db.log_export(st.session_state.user, project_id, "TXT")

    with col2:
        if st.download_button("📝 Markdown", f"# {title}\n\n{script_text}",
                               file_name=f"{title}_{timestamp}.md", use_container_width=True):
            db.log_export(st.session_state.user, project_id, "Markdown")

    with col3:
        export_json = {
            "title": title,
            "script": script_text,
            "camera_angles": db.safe_json_loads(project["camera_angles"], []) if project else st.session_state.get("camera_angles_data", []),
            "scene_breakdown": db.safe_json_loads(project["scene_breakdown"], []) if project else st.session_state.get("scene_breakdown_data", []),
            "storyboard": db.safe_json_loads(project["storyboard"], []) if project else st.session_state.get("storyboard_data", []),
            "characters": db.safe_json_loads(project["characters"], []) if project else st.session_state.get("characters_data", []),
        }
        if st.download_button("🗂️ JSON", json.dumps(export_json, indent=2),
                               file_name=f"{title}_{timestamp}.json", use_container_width=True):
            db.log_export(st.session_state.user, project_id, "JSON")

    with col4:
        try:
            pdf_buffer = _build_pdf(title, script_text)
            if st.download_button("📕 PDF", pdf_buffer, file_name=f"{title}_{timestamp}.pdf", use_container_width=True):
                db.log_export(st.session_state.user, project_id, "PDF")
        except Exception as e:
            st.error(f"PDF export failed: {e}")

    with col5:
        try:
            docx_buffer = _build_docx(title, script_text)
            if st.download_button("📘 DOCX", docx_buffer, file_name=f"{title}_{timestamp}.docx", use_container_width=True):
                db.log_export(st.session_state.user, project_id, "DOCX")
        except Exception as e:
            st.error(f"DOCX export failed: {e}")

    glass_close()

    st.caption("💡 JSON export includes Camera Angles, Scene Breakdown, Storyboard and "
               "Characters if you've generated them for this project.")
